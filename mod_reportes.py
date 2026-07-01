"""
Módulo: Reportes Radar
Rol: Exportación de informes oficiales en Word
"""
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from utils_c5i import COMUNAS_PURGADAS

def render_reportes(df_filtrado, f_inicio, f_fin):
    st.subheader("📄 Módulo de Exportación Oficial: Radar de Crisis (.docx)")
    
    if st.button("🚀 Compilar Informe Oficial", use_container_width=True, type="primary"):
        with st.spinner("Trazando gráficos en memoria e incrustando recursos visuales..."):
            try:
                fig_barras, ax_bar = plt.subplots(figsize=(7, 3.5))
                fig_barras.patch.set_facecolor('#ffffff')
                ax_bar.set_facecolor('#ffffff')
                df_tipos_rep = df_filtrado['tipologia_oficial'].value_counts() if not df_filtrado.empty else pd.Series()
                if not df_tipos_rep.empty:
                    df_tipos_rep.head(6).plot(kind='barh', color='#003366', ax=ax_bar)
                    ax_bar.set_title('Composición de Sucesos por Tipología', fontsize=11, fontweight='bold', color='#003366')
                    ax_bar.set_xlabel('Cantidad de Eventos', fontsize=9)
                    ax_bar.invert_yaxis()
                    plt.tight_layout()
                else:
                    ax_bar.text(0.5, 0.5, 'Sin masa crítica para graficar tipologías', ha='center', va='center')
                img_stream_barras = io.BytesIO()
                plt.savefig(img_stream_barras, format='png', dpi=200, bbox_inches='tight')
                img_stream_barras.seek(0)
                plt.close(fig_barras)

                fig_pie, ax_pie = plt.subplots(figsize=(5, 3.5))
                fig_pie.patch.set_facecolor('#ffffff')
                df_alertas_rep = df_filtrado['nivel_alerta'].value_counts() if not df_filtrado.empty and 'nivel_alerta' in df_filtrado.columns else pd.Series()
                colores_map = {'CRÍTICO': '#8B0000', 'ALTO': '#FF8C00', 'MEDIO': '#FFD700', 'BAJO': '#4682B4'}
                if not df_alertas_rep.empty:
                    cols_pie = [colores_map.get(x, '#808080') for x in df_alertas_rep.index]
                    df_alertas_rep.plot(kind='pie', autopct='%1.1f%%', colors=cols_pie, ax=ax_pie, startangle=90, textprops={'fontsize': 8})
                    ax_pie.set_ylabel('')
                    ax_pie.set_title('Distribución de Alertas', fontsize=11, fontweight='bold', color='#003366')
                    plt.tight_layout()
                else:
                    ax_pie.text(0.5, 0.5, 'Sin masa crítica', ha='center', va='center')
                img_stream_pie = io.BytesIO()
                plt.savefig(img_stream_pie, format='png', dpi=200, bbox_inches='tight')
                img_stream_pie.seek(0)
                plt.close(fig_pie)

                doc = Document()
                for section in doc.sections:
                    section.top_margin = Inches(0.8)
                    section.bottom_margin = Inches(0.8)
                    section.left_margin = Inches(0.8)
                    section.right_margin = Inches(0.8)

                style_normal = doc.styles['Normal']
                font = style_normal.font
                font.name = 'Arial'
                font.size = Pt(10.5)
                font.color.rgb = RGBColor(0x22, 0x22, 0x22)

                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_title = p_title.add_run("RADAR DE CRISIS - INFORME DE INTELIGENCIA TERRITORIAL\nGERENCIA DE PROTECCIÓN PATRIMONIAL")
                r_title.font.size = Pt(14)
                r_title.font.bold = True
                r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

                p_meta = doc.add_paragraph()
                p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_meta = p_meta.add_run(f"Confidencial - Estado Mayor CMPC | Fecha de Emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}\nVentana Analizada: {f_inicio.strftime('%d/%m/%Y')} al {f_fin.strftime('%d/%m/%Y')}")
                r_meta.font.size = Pt(9.5)
                r_meta.font.italic = True

                doc.add_paragraph()

                h1 = doc.add_heading("I. Apreciación Descriptiva y Contexto Territorial", level=1)
                h1.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

                total_ev = len(df_filtrado)
                crit_ev = len(df_filtrado[df_filtrado['nivel_alerta'] == 'CRÍTICO']) if total_ev > 0 and 'nivel_alerta' in df_filtrado.columns else 0
                ig_ev = len(df_filtrado[df_filtrado['es_rrss'] == True]) if total_ev > 0 and 'es_rrss' in df_filtrado.columns else 0
                prensa_ev = total_ev - ig_ev

                comunas_validas = []
                if total_ev > 0 and 'ubicacion' in df_filtrado.columns:
                    excluir_locs = ['no especificado', 'desconocido', 'sin dato', 'mzs', '', 'macrozona sur'] + COMUNAS_PURGADAS
                    comunas_serie = df_filtrado['ubicacion'].dropna().astype(str)
                    comunas_validas = comunas_serie[~comunas_serie.str.lower().str.strip().isin(excluir_locs)]

                comunas_afectadas = comunas_validas.nunique() if len(comunas_validas) > 0 else 0
                principales_comunas = ", ".join(comunas_validas.value_counts().head(3).index.tolist()) if len(comunas_validas) > 0 else "sectores focales del corredor"

                p_ap1 = doc.add_paragraph()
                p_ap1.paragraph_format.line_spacing = 1.15
                p_ap1.paragraph_format.space_after = Pt(6)
                p_ap1.add_run(
                    f"Durante el periodo sometido a auditoría, el sistema C5I procesó un total de {total_ev} eventos de interés "
                    f"operativo. La masa crítica destilada se compone de {prensa_ev} reportes extraídos desde partes de contingencia y prensa, "
                    f"más {ig_ev} trazas de inteligencia nativa interceptadas en redes sociales (Meta/Instagram). La conflictividad hostil "
                    f"y la presencia policial exhibieron una focalización que abarcó {comunas_afectadas} comunas territorialmente identificables, "
                    f"saturando de manera principal los ejes de {principales_comunas}."
                )

                p_ap2 = doc.add_paragraph()
                p_ap2.paragraph_format.line_spacing = 1.15
                p_ap2.paragraph_format.space_after = Pt(12)
                p_ap2.add_run(
                    f"El destilado algorítmico arroja que {crit_ev} sucesos asumen carácter CRÍTICO directo para la compañía "
                    f"al vulnerar o amenazar faenas silvícolas, maquinaria forestal o infraestructura patrimonial. Aquellos hitos "
                    f"vinculados a inversión comunitaria o pautas políticas públicas han sido filtrados para asegurar la objetividad "
                    f"de la presente matriz."
                )

                h_graf = doc.add_heading("II. Representación Gráfica de Métricas Operativas", level=1)
                h_graf.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

                p_g1 = doc.add_paragraph()
                p_g1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_g1.paragraph_format.space_after = Pt(4)
                r_g1_lbl = p_g1.add_run("Figura 1: Distribución por Tipología Operativa (Prensa e IG combinados)")
                r_g1_lbl.font.size = Pt(9.0)
                r_g1_lbl.font.italic = True

                doc.add_picture(img_stream_barras, width=Inches(5.8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()

                p_g2 = doc.add_paragraph()
                p_g2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_g2.paragraph_format.space_after = Pt(4)
                r_g2_lbl = p_g2.add_run("Figura 2: Proporción de Estados de Alerta en Ventana de Análisis")
                r_g2_lbl.font.size = Pt(9.0)
                r_g2_lbl.font.italic = True

                doc.add_picture(img_stream_pie, width=Inches(4.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()

                h2 = doc.add_heading("III. Detalle de Vulneraciones Críticas Directas", level=1)
                h2.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

                df_criticos = df_filtrado[df_filtrado['nivel_alerta'] == 'CRÍTICO'] if total_ev > 0 and 'nivel_alerta' in df_filtrado.columns else pd.DataFrame()

                if not df_criticos.empty:
                    table = doc.add_table(rows=1, cols=3)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Fecha'
                    hdr_cells[1].text = 'Comuna / Sector'
                    hdr_cells[2].text = 'Titular / Descripción Fáctica'
                    
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                    
                    for _, c_row in df_criticos.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(c_row.get('fecha_limpia', ''))
                        
                        loc_txt = str(c_row.get('ubicacion', 'MZS')).strip()
                        row_cells[1].text = loc_txt if loc_txt.lower() not in ['no especificado', 'desconocido'] else "Corredor Forestal"
                        
                        tit_txt = str(c_row.get('titular', ''))
                        act_txt = str(c_row.get('actor', 'N/A')).strip()
                        atrib = f" [Atribución: {act_txt}]" if act_txt.lower() not in ['desconocido', 'no atribuido', ''] else ""
                        row_cells[2].text = f"{tit_txt}{atrib}"
                        
                        for cell in row_cells:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(9.0)
                    
                    doc.add_paragraph()
                else:
                    p_safe = doc.add_paragraph()
                    p_safe.paragraph_format.space_after = Pt(12)
                    r_safe = p_safe.add_run("En la presente ventana analizada, la compuerta algorítmica no detectó sucesos directos de sabotaje contra el patrimonio o colaboradores de CMPC.")
                    r_safe.font.italic = True
                
                h_prosp = doc.add_heading("IV. Análisis Prospectivo y Escenarios de Riesgo", level=1)
                h_prosp.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                
                p_pr1 = doc.add_paragraph()
                p_pr1.paragraph_format.line_spacing = 1.15
                p_pr1.paragraph_format.space_after = Pt(6)
                p_pr1.add_run(
                    "Con base en la distribución espacial trazada y la evolución temporal capturada en los gráficos precedentes, "
                    "el sistema deduce una clara intencionalidad de sostenimiento asimétrico por parte de las orgánicas activas. "
                    "La presencia de pautas combinadas de allanamiento y respuestas armadas indica un nivel de fricción territorial "
                    "elevado que tiende a desplazar el riesgo logístico hacia corredores secundarios de transporte forestal."
                )
                
                p_pr2 = doc.add_paragraph()
                p_pr2.paragraph_format.line_spacing = 1.15
                p_pr2.paragraph_format.space_after = Pt(12)
                p_pr2.add_run(
                    "Se proyecta que las instalaciones industriales y anillos perimetrales mantengan un estatus operativo estable "
                    "siempre y cuando se garantice la retroalimentación continua del perímetro de Geofencing y se apliquen las restricciones "
                    "de convoyes nocturnos en los tramos críticos de Arauco y Malleco."
                )
                
                h3 = doc.add_heading("V. Directrices de Mando Permanentes", level=1)
                h3.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                
                directrices = [
                    "Sostener la inmovilización nocturna para convoyes de carga en rutas aledañas a los sectores con registros críticos.",
                    "Mantener sincronizados los avisos preventivos entre los monitores de plataforma y jefaturas de zona.",
                    "Actualizar semanalmente las coordenadas de faena activa en la base central para asegurar la calibración del Geofencing."
                ]
                
                for idx, d_txt in enumerate(directrices, 1):
                    p_dir = doc.add_paragraph()
                    p_dir.paragraph_format.left_indent = Inches(0.2)
                    p_dir.paragraph_format.space_after = Pt(4)
                    p_dir.add_run(f"{idx}. ").font.bold = True
                    p_dir.add_run(d_txt)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.success("✔️ Reporte 'Radar de Crisis' compilado con éxito (Gráficos nativos incrustados).")
                st.download_button(
                    label="📥 Descargar Documento Oficial (.docx)",
                    data=buffer,
                    file_name=f"Radar_de_Crisis_CMPC_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e_doc:
                st.error(f"Error interno al destilar el documento Word con gráficos: {e_doc}")