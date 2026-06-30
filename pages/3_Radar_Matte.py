import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import yfinance as yf
import feedparser
import urllib.parse
import urllib.request
import requests
import os
import io
import time
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="Radar de Inteligencia: Grupo Matte", layout="wide")

ARCHIVO_HISTORIAL = "historial_noticias.csv"

# [PEGA AQUÍ TU ENLACE DE MAKE.COM PARA TELEGRAM]
URL_WEBHOOK_MAKE = "https://hook.us2.make.com/xit3mte8pkkyibhwaz382ns909edm4ht"

# BARRERA ANTI-RUIDO (Análisis léxico estricto mediante bigramas y trigramas)
PALABRAS_CLAVE_CRITICAS = [
    "ataque incendiario", "ataque armado", "incendio forestal", "multa ambiental", 
    "demanda colectiva", "huelga legal", "toma de terreno", "mininco camionero", "usurpación violenta", 
    "fallo judicial", "querella criminal", "fraude corporativo", "estafa financiera", 
    "atentado explosivo", "emboscada armada", "sabotaje industrial", "enfrentamiento armado", 
    "corte de ruta", "intercambio de disparos", "impactos de bala", "coordinadora arauco malleco",
    "resistencia mapuche malleco", "resistencia mapuche lafkenche", "weichan auka mapu", 
    "organización radical", "grupo armado", "reivindicación territorial", "bloqueo de ruta",
    "fiscalía nacional económica", "tribunal de defensa", "servicio nacional del consumidor"
]

# FRENTES DE BÚSQUEDA
FRENTES = [
    {"titulo": "🏢 Cúpula y Familia", "query": '"Familia Matte" OR "Grupo Matte" OR "Minera Valparaíso"'},
    {"titulo": "⚡ Colbún: Socio-Ambiental", "query": '"Colbún" AND (agua OR conflicto OR comunidad OR huelga)'},
    {"titulo": "📱 Entel: Regulatorio y Quejas", "query": '"Entel" AND (multa OR TDLC OR Subtel OR SERNAC OR reclamo)'},
    {"titulo": "🌲 CMPC / Mininco: Macrozona Sur", "query": '("CMPC" OR "Mininco" OR "Forestal Mininco") AND (Mapuche OR Macrozona OR ataque OR forestal OR atentado OR Temucuicui OR weichafe OR camionero OR sabotaje OR "fuera forestales")'},
    {"titulo": "🌐 CMPC / Mininco: Internacional", "query": '("CMPC" OR "Mininco") AND (Brasil OR Perú OR México OR Guaíba OR conflicto)'},
    {"titulo": "🏦 BICE: Reclamos y Finanzas", "query": '"Banco BICE" OR "BICE Vida" AND (reclamo OR SERNAC OR estafa OR fraude)'}
]

COLORES_CORPORATIVOS = {
    'Colbún': '#FF4500',
    'Entel': '#1E90FF',
    'CMPC': '#00A86B',
    'Minera Valparaíso (Holding)': '#FFD700',
    'Volcán': '#808080',
    'BICE': '#4B0082'
}

# --- MOTORES DE EXTRACCIÓN ---

@st.cache_data(ttl=900) 
def obtener_datos_bolsa():
    tickers = {
        "Colbún": "COLBUN.SN", 
        "Entel": "ENTEL.SN", 
        "CMPC": "CMPC.SN",
        "Minera Valparaíso (Holding)": "MINERA.SN"
    }
    df_list = []
    kpis = {} 
    
    for nombre, ticker in tickers.items():
        try:
            empresa_data = yf.Ticker(ticker)
            historia = empresa_data.history(period="3mo")
            if not historia.empty:
                historia['Empresa'] = nombre
                historia['Fecha'] = historia.index
                
                ultimo_precio = historia['Close'].iloc[-1]
                precio_anterior = historia['Close'].iloc[-2] if len(historia) > 1 else ultimo_precio
                variacion = ((ultimo_precio - precio_anterior) / precio_anterior) * 100
                kpis[nombre] = {"precio": ultimo_precio, "variacion": variacion}
                
                historia = historia[['Fecha', 'Open', 'High', 'Low', 'Close', 'Volume', 'Empresa']]
                df_list.append(historia)
        except Exception:
            pass 
            
    if df_list:
        return pd.concat(df_list), kpis
    return pd.DataFrame(), {}

def buscar_noticias_en_vivo(query_booleana):
    query_codificada = urllib.parse.quote(query_booleana)
    url = f"https://news.google.com/rss/search?q={query_codificada}+when:12h&hl=es-419&gl=CL&ceid=CL:es-419"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0 Safari/537.36'}
    try:
        respuesta = requests.get(url, headers=headers, timeout=10)
        feed = feedparser.parse(respuesta.content)
        return feed.entries
    except Exception:
        return []

# [NUEVO]: MOTOR ACORTADOR DE URL
def acortar_url(url_larga):
    try:
        api_url = f"http://tinyurl.com/api-create.php?url={urllib.parse.quote(url_larga)}"
        respuesta = urllib.request.urlopen(api_url, timeout=5)
        return respuesta.read().decode('utf-8')
    except Exception:
        return url_larga # Fallback: si el acortador falla, envía el original para no perder la alerta

def enviar_alerta_webhook(frente, titulo, link, fecha_publicacion):
    if URL_WEBHOOK_MAKE != "":
        datos_alerta = {
            "frente": frente,
            "titulo": titulo,
            "link": link,
            "fecha": fecha_publicacion
        }
        try:
            requests.post(URL_WEBHOOK_MAKE, json=datos_alerta, timeout=5)
        except Exception:
            pass 

def ejecutar_aspiradora_noticias():
    if os.path.exists(ARCHIVO_HISTORIAL):
        df_historial = pd.read_csv(ARCHIVO_HISTORIAL)
    else:
        df_historial = pd.DataFrame(columns=["Fecha_Captura", "Frente", "Titulo", "Link", "Fecha_Publicacion"])
    
    nuevos_registros = []
    links_guardados = set(df_historial["Link"].tolist())
    limite_tiempo = datetime.now() - timedelta(hours=12)
    
    for frente in FRENTES:
        noticias_vivas = buscar_noticias_en_vivo(frente["query"])
        for noticia in noticias_vivas:
            if noticia.link not in links_guardados:
                
                if hasattr(noticia, 'published_parsed') and noticia.published_parsed:
                    fecha_matematica = datetime.fromtimestamp(time.mktime(noticia.published_parsed))
                    if fecha_matematica < limite_tiempo:
                        continue 
                
                nuevos_registros.append({
                    "Fecha_Captura": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Frente": frente["titulo"],
                    "Titulo": noticia.title,
                    "Link": noticia.link,
                    "Fecha_Publicacion": noticia.published
                })
                links_guardados.add(noticia.link)
                
                titulo_minusculas = noticia.title.lower()
                if any(palabra in titulo_minusculas for palabra in PALABRAS_CLAVE_CRITICAS):
                    # [MODIFICADO]: Acortamos la URL antes de enviarla a Telegram
                    link_limpio = acortar_url(noticia.link)
                    enviar_alerta_webhook(frente["titulo"], noticia.title, link_limpio, noticia.published)
                
    if nuevos_registros:
        df_nuevos = pd.DataFrame(nuevos_registros)
        df_consolidado = pd.concat([df_historial, df_nuevos], ignore_index=True)
        df_consolidado.to_csv(ARCHIVO_HISTORIAL, index=False)
        return len(nuevos_registros)
    return 0

def crear_grafico_estetico(df, empresa, color_hex):
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=df['Fecha'], y=df['Close'],
        mode='lines',
        name=empresa,
        line=dict(color=color_hex, width=3)
    ))
    fig.update_layout(
        title=dict(text=f"Cotización {empresa} (Últimos 3 Meses)", font=dict(size=18)),
        margin=dict(l=10, r=10, t=40, b=10),
        plot_bgcolor="rgba(0,0,0,0)",
        paper_bgcolor="rgba(0,0,0,0)",
        xaxis=dict(showgrid=False, zeroline=False),
        yaxis=dict(showgrid=False, zeroline=False, autorange=True)
    )
    return fig

# --- MOTOR DE GENERACIÓN DE DOCUMENTO WORD (FLASH REPORT) ---
def generar_reporte_word(df_alertas, fig_radar_export, fig_norm_export):
    doc = Document()
    
    tit_principal = doc.add_heading('FLASH REPORT - INTELIGENCIA MULTISECTORIAL', 0)
    tit_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtit = doc.add_paragraph('HOLDING MATTE (MINERA VALPARAÍSO S.A.)')
    subtit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Confidencial - Plataforma de Monitoreo OSINT\nFecha de Emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
    
    doc.add_heading('I. Apreciación Descriptiva y Contexto Corporativo', level=1)
    doc.add_paragraph(f"El presente informe ejecutivo sintetiza el estado de situación de las filiales del Holding Matte. Durante el periodo de análisis, el sistema ha procesado {len(df_alertas)} trazas totales de exposición a riesgo en fuentes abiertas.")
    doc.add_paragraph("A diferencia de matrices pasadas enfocadas exclusivamente en la fricción forestal, el actual ecosistema de riesgos exige un monitoreo transversal: presión regulatoria y de mercado en telecomunicaciones (Entel), tensión hídrico-comunitaria en la matriz energética (Colbún), resguardo del estándar ESG en el sector financiero (BICE), y la protección patrimonial en la Macrozona Sur e internacional (CMPC).")
    
    doc.add_heading('II. Mapeo Visual de Riesgos y Mercado', level=1)
    try:
        img_radar = io.BytesIO()
        fig_radar_export.update_layout(
            paper_bgcolor="white", 
            font=dict(color="black"), 
            polar=dict(bgcolor="white", radialaxis=dict(gridcolor="lightgrey", linecolor="black"))
        )
        fig_radar_export.write_image(img_radar, format="png", width=600, height=400)
        img_radar.seek(0)
        doc.add_picture(img_radar, width=Inches(6))
        
        if fig_norm_export:
            img_norm = io.BytesIO()
            fig_norm_export.update_layout(
                paper_bgcolor="white", 
                plot_bgcolor="white", 
                font=dict(color="black"), 
                title=dict(text="Rendimiento Comparado Base 0%", font=dict(color="black")),
                xaxis=dict(showgrid=False, zeroline=False, linecolor="black", tickfont=dict(color="black")),
                yaxis=dict(showgrid=True, gridcolor="lightgrey", zeroline=True, zerolinecolor="grey", linecolor="black", tickfont=dict(color="black")),
                legend=dict(font=dict(color="black"))
            )
            fig_norm_export.write_image(img_norm, format="png", width=600, height=350)
            img_norm.seek(0)
            doc.add_picture(img_norm, width=Inches(6))
    except Exception as e:
        doc.add_paragraph("(Aviso del Sistema: No se pudieron incrustar los gráficos visuales. Asegúrese de que la librería 'kaleido' esté instalada).")
    
    doc.add_heading('III. Destilado Crítico: Top 3 Alertas por Frente', level=1)
    
    frentes_unicos = [f["titulo"] for f in FRENTES] 
    for frente in frentes_unicos:
        doc.add_heading(f"{frente.replace('titulo: ', '')}", level=2)
        df_frente = df_alertas[df_alertas['Frente'] == frente].sort_values(by="Fecha_Captura", ascending=False).head(3)
        
        if not df_frente.empty:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Fecha Pub.'
            hdr_cells[1].text = 'Evento Detectado'
            
            hdr_cells[0].width = Inches(1.5)
            hdr_cells[1].width = Inches(4.5)
            
            for index, row in df_frente.iterrows():
                row_cells = table.add_row().cells
                fecha_corta = str(row['Fecha_Publicacion'])[:25] if pd.notna(row['Fecha_Publicacion']) else "Reciente"
                row_cells[0].text = fecha_corta
                row_cells[1].text = row['Titulo']
        else:
            doc.add_paragraph("Sin alertas críticas detectadas en el periodo reciente.")
            
    doc.add_heading('IV. Análisis Prospectivo Multi-Sectorial', level=1)
    doc.add_paragraph("1. Sector Forestal (CMPC): La conflictividad asimétrica obliga a mantener un estricto control logístico. La exposición internacional (Guaíba) añade un vector secundario de riesgo bursátil.")
    doc.add_paragraph("2. Sector Energía (Colbún): Los desafíos de transición ecológica y derechos de agua requieren un enfoque preventivo en el relacionamiento comunitario para evitar paralizaciones de nuevos proyectos.")
    doc.add_paragraph("3. Sector Telecomunicaciones (Entel): Se observa que la matriz de riesgo está fuertemente correlacionada a las quejas ciudadanas masivas y litigios antimonopolio (TDLC/Subtel).")
    doc.add_paragraph("4. Sector Financiero (BICE): Es imperativo vigilar los riesgos reputacionales en el financiamiento corporativo para evitar el contagio hacia la banca privada del holding.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

noticias_nuevas_capturadas = ejecutar_aspiradora_noticias()

# --- GENERACIÓN DE GRÁFICOS GLOBALES PARA EL DASHBOARD Y EL WORD ---
datos_bolsa, kpis_bolsa = obtener_datos_bolsa()

fig_norm = None
if not datos_bolsa.empty:
    df_norm = datos_bolsa.copy()
    df_norm['Variacion_Base'] = df_norm.groupby('Empresa')['Close'].transform(lambda x: ((x / x.iloc[0]) - 1) * 100)
    fig_norm = px.line(df_norm, x='Fecha', y='Variacion_Base', color='Empresa', color_discrete_map=COLORES_CORPORATIVOS, title="Rendimiento Comparativo Normalizado (Base 0%)")
    fig_norm.update_layout(
        margin=dict(l=10, r=10, t=40, b=10),
        plot_bgcolor="rgba(0,0,0,0)",
        paper_bgcolor="rgba(0,0,0,0)",
        xaxis=dict(showgrid=False, zeroline=False),
        yaxis=dict(showgrid=True, gridcolor="rgba(128,128,128,0.2)", zeroline=True, zerolinecolor="rgba(255,255,255,0.5)")
    )

df_riesgos_radar = pd.DataFrame({
    'Empresa': ['CMPC', 'Colbún', 'Entel', 'Volcán', 'BICE'],
    'Riesgo': [9.5, 8.0, 7.0, 5.5, 4.0]
})
fig_radar = go.Figure(data=go.Scatterpolar(
    r=df_riesgos_radar['Riesgo'],
    theta=df_riesgos_radar['Empresa'],
    fill='toself',
    line_color='#FF4500',
    fillcolor='rgba(255, 69, 0, 0.3)'
))
fig_radar.update_layout(
    polar=dict(radialaxis=dict(visible=True, range=[0, 10], gridcolor="rgba(128,128,128,0.2)"), bgcolor="rgba(0,0,0,0)"),
    showlegend=False,
    margin=dict(t=10, l=10, r=10, b=10),
    paper_bgcolor="rgba(0,0,0,0)"
)

# --- INTERFAZ DEL DASHBOARD ---

st.title("🕵️‍♂️ Central de Inteligencia: Grupo Matte")

if noticias_nuevas_capturadas > 0:
    st.toast(f"🤖 Aspiradora activa: {noticias_nuevas_capturadas} nuevas incidencias archivadas.", icon="📥")

tab_global, tab_cmpc, tab_colbun, tab_entel, tab_bice, tab_tweetdeck, tab_archivo = st.tabs([
    "🏛️ Comando Central (Global)",
    "🌲 CMPC / Mininco",
    "⚡ Colbún",
    "📱 Entel",
    "🏦 BICE / Volcán",
    "📡 TweetDeck",
    "🗄️ Archivo Histórico y Reportes"
])

configuracion_grafico = {'displayModeBar': False}

with tab_global:
    st.header("Radiografía de la Matriz y Ecosistema Matte")
    if kpis_bolsa:
        columnas_kpi = st.columns(len(kpis_bolsa))
        for i, (empresa, datos) in enumerate(kpis_bolsa.items()):
            columnas_kpi[i].metric(label=f"Acción: {empresa}", value=f"${datos['precio']:,.0f}", delta=f"{datos['variacion']:.2f}%")
        st.divider()

        col_g1, col_g2 = st.columns(2)
        with col_g1:
            if not datos_bolsa.empty and "Minera Valparaíso (Holding)" in kpis_bolsa:
                df_holding = datos_bolsa[datos_bolsa['Empresa'] == 'Minera Valparaíso (Holding)']
                fig_holding = crear_grafico_estetico(df_holding, "Minera Valparaíso (Holding)", COLORES_CORPORATIVOS["Minera Valparaíso (Holding)"])
                st.plotly_chart(fig_holding, use_container_width=True, config=configuracion_grafico)
                
        with col_g2:
            st.markdown("### 📊 Rendimiento Comparativo")
            st.caption("Muestra quién tracciona o lastra al holding comparando el crecimiento porcentual.")
            if fig_norm:
                st.plotly_chart(fig_norm, use_container_width=True, config=configuracion_grafico)

        col_g3, col_g4 = st.columns(2)
        with col_g3:
            st.markdown("### 🧩 Composición del Imperio (Estimación)")
            df_tree = pd.DataFrame({
                'Sector': ['Forestal e Industrial', 'Energía', 'Telecomunicaciones', 'Financiero', 'Construcción'],
                'Empresa': ['CMPC', 'Colbún', 'Entel', 'BICECORP', 'Volcán'],
                'Peso': [45, 25, 15, 10, 5]
            })
            fig_tree = px.treemap(df_tree, path=['Sector', 'Empresa'], values='Peso', color='Peso', color_continuous_scale='YlOrBr')
            fig_tree.update_layout(margin=dict(t=10, l=10, r=10, b=10), paper_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_tree, use_container_width=True, config=configuracion_grafico)

        with col_g4:
            st.markdown("### ⚠️ Radar Térmico de Conflictividad")
            st.plotly_chart(fig_radar, use_container_width=True, config=configuracion_grafico)

with tab_cmpc:
    st.header("Frente Forestal: CMPC y Mininco")
    col_grafico, col_noticias = st.columns([2, 1])
    with col_grafico:
        if not datos_bolsa.empty and "CMPC" in kpis_bolsa:
            df_cmpc = datos_bolsa[datos_bolsa['Empresa'] == 'CMPC']
            fig_cmpc = crear_grafico_estetico(df_cmpc, "CMPC", COLORES_CORPORATIVOS["CMPC"])
            st.plotly_chart(fig_cmpc, use_container_width=True, config=configuracion_grafico)
    with col_noticias:
        st.subheader("🚨 Radar Macrozona Sur")
        n_macrozona = buscar_noticias_en_vivo(FRENTES[3]["query"])
        if n_macrozona:
            for n in n_macrozona[:4]:
                st.markdown(f"🔸 [{n.title}]({n.link})")

with tab_colbun:
    st.header("Frente Energético: Colbún S.A.")
    col_grafico, col_noticias = st.columns([2, 1])
    with col_grafico:
        if not datos_bolsa.empty and "Colbún" in kpis_bolsa:
            df_colb = datos_bolsa[datos_bolsa['Empresa'] == 'Colbún']
            fig_colb = crear_grafico_estetico(df_colb, "Colbún", COLORES_CORPORATIVOS["Colbún"])
            st.plotly_chart(fig_colb, use_container_width=True, config=configuracion_grafico)
    with col_noticias:
        st.subheader("🚨 Conflictos Hídricos")
        n_colbun = buscar_noticias_en_vivo(FRENTES[1]["query"])
        if n_colbun:
            for n in n_colbun[:4]:
                st.markdown(f"🔸 [{n.title}]({n.link})")

with tab_entel:
    st.header("Frente Tecnológico: Entel")
    col_grafico, col_noticias = st.columns([2, 1])
    with col_grafico:
        if not datos_bolsa.empty and "Entel" in kpis_bolsa:
            df_entel = datos_bolsa[datos_bolsa['Empresa'] == 'Entel']
            fig_entel = crear_grafico_estetico(df_entel, "Entel", COLORES_CORPORATIVOS["Entel"])
            st.plotly_chart(fig_entel, use_container_width=True, config=configuracion_grafico)
    with col_noticias:
        st.subheader("📉 Radar Regulación")
        n_entel = buscar_noticias_en_vivo(FRENTES[2]["query"])
        if n_entel:
            for n in n_entel[:4]:
                st.markdown(f"🔸 [{n.title}]({n.link})")

with tab_bice:
    st.header("Banca e Industria: BICECORP y Volcán")
    col_b1, col_b2 = st.columns(2)
    with col_b1:
        st.subheader("🏦 Banco BICE y BICE Vida")
        n_bice = buscar_noticias_en_vivo(FRENTES[5]["query"])
        if n_bice:
            for n in n_bice[:4]:
                st.markdown(f"🔸 [{n.title}]({n.link})")
    with col_b2:
        st.subheader("🏗️ Volcán")
        st.warning("**Conflicto Histórico:** Minería de yeso en Cajón del Maipo.")

with tab_tweetdeck:
    st.header("Monitor Consolidado (TweetDeck)")
    cols = st.columns(4)
    for i in range(4):
        with cols[i]:
            st.subheader(FRENTES[i]["titulo"])
            st.divider() 
            resultados = buscar_noticias_en_vivo(FRENTES[i]["query"])
            if resultados:
                with st.container(height=400):
                    for entrada in resultados[:10]: 
                        st.markdown(f"**[{entrada.title}]({entrada.link})**")
                        st.write("---")

with tab_archivo:
    col_izq, col_der = st.columns([2, 1])
    with col_izq:
        st.header("🗄️ Repositorio de Base de Datos")
        if os.path.exists(ARCHIVO_HISTORIAL):
            df_archivo_data = pd.read_csv(ARCHIVO_HISTORIAL)
            filtro_frente = st.selectbox("Filtrar archivo por Frente:", ["Todos"] + df_archivo_data["Frente"].unique().tolist())
            df_filtrado = df_archivo_data if filtro_frente == "Todos" else df_archivo_data[df_archivo_data["Frente"] == filtro_frente]
            st.dataframe(df_filtrado[["Fecha_Captura", "Frente", "Titulo", "Fecha_Publicacion"]].sort_values(by="Fecha_Captura", ascending=False), use_container_width=True)
        else:
            st.info("El archivo histórico está vacío. Espera a capturar alertas.")
            
    with col_der:
        st.header("📄 Exportar Inteligencia")
        st.markdown("Genera el reporte ejecutivo. El sistema incrustará los gráficos operativos y el **Top 3 de noticias más recientes** por cada empresa.")
        
        if os.path.exists(ARCHIVO_HISTORIAL):
            archivo_word = generar_reporte_word(df_archivo_data, fig_radar, fig_norm)
            st.download_button(
                label="📥 Generar Reporte Flash (.docx)",
                data=archivo_word,
                file_name=f"Flash_Report_Matte_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
